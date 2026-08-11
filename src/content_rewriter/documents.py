import re
from enum import Enum
from pathlib import Path

import docx

PLAIN_SUFFIXES = {".txt", ""}
MARKDOWN_SUFFIXES = {".md", ".markdown", ".mdown", ".mkd"}
DOCX_SUFFIXES = {".docx"}
TEXT_SUFFIXES = PLAIN_SUFFIXES | MARKDOWN_SUFFIXES
READABLE = TEXT_SUFFIXES | DOCX_SUFFIXES

HEADING_LEVEL = re.compile(r"^Heading (\d+)")
MARKDOWN_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
MARKDOWN_BULLET = re.compile(r"^\s*[-*+]\s+(.*)$")
MARKDOWN_NUMBER = re.compile(r"^\s*\d+[.)]\s+(.*)$")
MARKDOWN_QUOTE = re.compile(r"^\s*>\s?(.*)$")
INLINE_SPAN = re.compile(r"(\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*[^*\n]+?\*|`[^`\n]+?`)")

STYLE_BULLET = "List Bullet"
STYLE_NUMBER = "List Number"
STYLE_QUOTE = "Quote"


class Kind(str, Enum):
    PLAIN = "plain"
    MARKDOWN = "markdown"


class UnsupportedFormat(Exception):
    pass


def supported_filter() -> str:
    return (
        "Documents (*.txt *.md *.markdown *.docx);;"
        "Text (*.txt *.md *.markdown);;"
        "Word (*.docx);;"
        "All files (*)"
    )


def kind_of(path) -> Kind:
    if path is None:
        return Kind.PLAIN
    suffix = Path(path).suffix.lower()
    if suffix in MARKDOWN_SUFFIXES or suffix in DOCX_SUFFIXES:
        return Kind.MARKDOWN
    return Kind.PLAIN


def load(path: Path) -> str:
    path = Path(path)
    if not path.is_file():
        raise FileNotFoundError(f"No such file: {path}")

    suffix = path.suffix.lower()
    if suffix in TEXT_SUFFIXES:
        return path.read_text(encoding="utf-8", errors="replace")
    if suffix in DOCX_SUFFIXES:
        return _docx_to_markdown(docx.Document(path))

    raise UnsupportedFormat(f"{suffix or path.name} is not supported. Use .txt, .md or .docx.")


def save(path: Path, text: str) -> Path:
    path = Path(path)
    suffix = path.suffix.lower()

    if suffix in DOCX_SUFFIXES:
        _markdown_to_docx(text).save(path)
        return path

    if suffix in TEXT_SUFFIXES:
        path.write_text(text, encoding="utf-8")
        return path

    raise UnsupportedFormat(f"Cannot write {suffix or path.name}.")


def output_path_for(path: Path) -> Path:
    path = Path(path)
    return path.with_name(f"{path.stem}.rewritten{path.suffix.lower() or '.txt'}")


def _docx_to_markdown(document) -> str:
    blocks = []
    for paragraph in document.paragraphs:
        text = _runs_to_markdown(paragraph).strip()
        if not text:
            continue

        style = (paragraph.style.name or "").strip()
        heading = HEADING_LEVEL.match(style)

        if heading:
            blocks.append((None, f"{'#' * min(int(heading.group(1)), 6)} {text}"))
        elif style == "Title":
            blocks.append((None, f"# {text}"))
        elif STYLE_BULLET in style:
            blocks.append(("bullet", f"- {text}"))
        elif STYLE_NUMBER in style:
            blocks.append(("number", f"1. {text}"))
        elif STYLE_QUOTE in style:
            blocks.append((None, f"> {text}"))
        else:
            blocks.append((None, text))

    rendered = ""
    previous = None
    for index, (listing, line) in enumerate(blocks):
        if index:
            rendered += "\n" if listing is not None and listing == previous else "\n\n"
        rendered += line
        previous = listing
    return rendered


def _runs_to_markdown(paragraph) -> str:
    pieces = []
    for run in paragraph.runs:
        text = run.text
        if not text.strip():
            pieces.append(text)
            continue

        leading = text[: len(text) - len(text.lstrip())]
        trailing = text[len(text.rstrip()) :]
        core = text.strip()

        if run.bold and run.italic:
            core = f"***{core}***"
        elif run.bold:
            core = f"**{core}**"
        elif run.italic:
            core = f"*{core}*"

        pieces.append(f"{leading}{core}{trailing}")
    return "".join(pieces) or paragraph.text


def _markdown_to_docx(text: str):
    document = docx.Document()
    buffer = []

    def flush():
        if buffer:
            _add_runs(document.add_paragraph(), " ".join(buffer))
            buffer.clear()

    for block in re.split(r"\n[ \t]*\n", text):
        for line in block.splitlines():
            stripped = line.strip()
            if not stripped:
                continue

            heading = MARKDOWN_HEADING.match(stripped)
            bullet = MARKDOWN_BULLET.match(stripped)
            number = MARKDOWN_NUMBER.match(stripped)
            quote = MARKDOWN_QUOTE.match(stripped)

            if heading:
                flush()
                level = min(len(heading.group(1)), 9)
                _add_runs(document.add_paragraph(style=f"Heading {level}"), heading.group(2))
            elif number:
                flush()
                _add_runs(document.add_paragraph(style=STYLE_NUMBER), number.group(1))
            elif bullet:
                flush()
                _add_runs(document.add_paragraph(style=STYLE_BULLET), bullet.group(1))
            elif quote:
                flush()
                _add_runs(document.add_paragraph(style=STYLE_QUOTE), quote.group(1))
            else:
                buffer.append(stripped)
        flush()

    return document


def _add_runs(paragraph, text: str):
    for piece in INLINE_SPAN.split(text):
        if not piece:
            continue

        if piece.startswith("***") and piece.endswith("***"):
            run = paragraph.add_run(piece[3:-3])
            run.bold = True
            run.italic = True
        elif piece.startswith("**") and piece.endswith("**"):
            paragraph.add_run(piece[2:-2]).bold = True
        elif piece.startswith("*") and piece.endswith("*"):
            paragraph.add_run(piece[1:-1]).italic = True
        elif piece.startswith("`") and piece.endswith("`"):
            run = paragraph.add_run(piece[1:-1])
            run.font.name = "Courier New"
        else:
            paragraph.add_run(piece)
    return paragraph
